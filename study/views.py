import base64
import docx2txt
import random
import openpyxl
import pandas as pd
import pdb;

from django.http import HttpResponse
from django.http import HttpResponse
from django.conf import settings
from django.shortcuts import render
from django.utils.translation import gettext as _

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from io import BytesIO
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import Alignment
from openpyxl.styles import Font
from openpyxl_image_loader import SheetImageLoader
from PIL import Image as PILImage

from .models import Question, Exam
from .models import SUBJECT_CHOICES, CATEGORY_CHOICES, QUESTION_IMAGE_PATH
from .utils import get_question_verbose_name, get_choice_key_by_value, arabic_numerals_to_chinese_numerals, add_hyperlink

IMAGE_COLUMN = 'F'

def index(request):
    messages = []
    if request.method == 'POST':
        uploaded_file = ''
        try:
            uploaded_file = request.FILES['excel_file']
        except:
            messages.append(f"{_('没有上传对应文件')}")
            return render(request, 'index.html', {'messages': messages})

        if uploaded_file.name.endswith('.xlsx'):
            df = pd.read_excel(uploaded_file, engine='openpyxl', na_filter = False)

            #loading the Excel File and the sheet
            pxl_doc = openpyxl.load_workbook(uploaded_file)
            sheet = pxl_doc['Sheet']

            #calling the image_loader
            image_loader = SheetImageLoader(sheet)

            for index, row in df.iterrows():
                id = row['id']
                grade = row[get_question_verbose_name('grade')]
                subject = get_choice_key_by_value(SUBJECT_CHOICES, row[get_question_verbose_name('subject')]) # value to key
                title = row[get_question_verbose_name('title')]
                description_above_image = row[get_question_verbose_name('description_above_image')]
                image_path = ''
                try:
                    image = image_loader.get(f'{IMAGE_COLUMN}{index + 2}')
                    image_path = f'{QUESTION_IMAGE_PATH}/{id}-{title}.png'
                    image.save(f'media/${image_path}')
                except Exception as e:
                    print(f'save image failed: {e}')
                description_below_image = row[get_question_verbose_name('description_below_image')]
                classroom_exercises = row[get_question_verbose_name('classroom_exercises')]
                score = row[get_question_verbose_name('number_of_errors')]
                category = get_choice_key_by_value(CATEGORY_CHOICES, row[get_question_verbose_name('category')]) # value to key
                fixed = row[get_question_verbose_name('fixed')]
                exam_name = row[get_question_verbose_name('exam_name')]
                exam_times = row[get_question_verbose_name('exam_times')]
                md5_value = row[get_question_verbose_name('md5_value')]
                answer = row[get_question_verbose_name('answer')]
                released = row[get_question_verbose_name('released')]
                creator = row[get_question_verbose_name('creator')]

                if Question.objects.filter(id=id) or Question.objects.filter(md5_value=md5_value):
                    if len(messages) < 100:
                        messages.append(f"Line {index + 2}: {_('重复数据无法导入')}")
                    elif len(messages) < 101:
                        messages.append(f"......")
                    continue

                empty_field = ''
                if not grade:
                    empty_field = 'grade'
                elif not subject:
                    empty_field = 'subject'
                elif not score:
                    empty_field = 'score'
                elif not category:
                    empty_field = 'category'
                
                if empty_field:
                    messages.clear()
                    messages.append(f"Line {index + 2}: \"{get_question_verbose_name(empty_field)}\" {_('不能为空')}")
                    break
                else:
                    question = Question.objects.create(
                        grade=grade,
                        subject=subject,
                        title=title,
                        description_above_image=description_above_image,
                        image=image_path,
                        description_below_image=description_below_image,
                        classroom_exercises=classroom_exercises,
                        score=score,
                        category=category,
                        fixed=fixed,
                        exam_name=exam_name,
                        exam_times=exam_times,
                        md5_value=md5_value,
                        answer=answer,
                        released=released,
                        creator=creator
                    )

                    if len(messages) < 10:
                        messages.append(f"Line {index + 2}: {_('数据已成功导入！')}")
                    elif len(messages) < 11:
                        messages.append(f"......")
        else:
            messages = f"_('文件格式不正确')"
    
    subjects = []
    for subject in SUBJECT_CHOICES:
        subjects.append(subject[1])

    return render(request, 'index.html', {'messages': messages, 'subjects': subjects})


def export_to_excel(request):
    queryset = Question.objects.all()

    # Create an Excel workbook
    workbook = Workbook()
    sheet = workbook.active

    # Prepare data (as a list of dictionaries)
    data = [
                {
                    "id": item.id, get_question_verbose_name('grade'): item.grade, get_question_verbose_name('subject'): item.subject, \
                    get_question_verbose_name('title'): item.title, get_question_verbose_name('description_above_image'): item.description_above_image, \
                    get_question_verbose_name('image'): item.image, get_question_verbose_name('description_below_image'): item.description_below_image, \
                    get_question_verbose_name('classroom_exercises'): item.classroom_exercises, get_question_verbose_name('score'): item.score, \
                    get_question_verbose_name('number_of_errors'): item.number_of_errors, get_question_verbose_name('category'): item.category, \
                    get_question_verbose_name('fixed'): item.fixed, get_question_verbose_name('exam_name'): item.exam_name, \
                    get_question_verbose_name('exam_times'): item.exam_times, get_question_verbose_name('md5_value'): item.md5_value, \
                    get_question_verbose_name('answer'): item.answer, \
                    get_question_verbose_name('released'): item.released, get_question_verbose_name('creator'): item.creator, \
                    get_question_verbose_name('created'): item.created, get_question_verbose_name('modified'): item.modified
                } for item in queryset
            ]
    
    # Add headers and set font style
    sheet.append(list(data[0]))
    # Create a Font object with the desired size and bold settings
    font = Font(size=16, bold=True)
    # Iterate through the cells in the specified row and apply the style
    for cell in sheet[1]:
        cell.font = font

    # Process and insert images into Excel
    for row, item in enumerate(data, start=2):
        # id
        sheet.cell(row=row, column=1, value=item[list(item)[0]])
        # grade
        sheet.cell(row=row, column=2, value=item[list(item)[1]])
        # subject
        sheet.cell(row=row, column=3, value=SUBJECT_CHOICES[item[list(item)[2]] - 1][1]) # Change to choice value
        # title
        sheet.cell(row=row, column=4, value=item[list(item)[3]])
        # description_above_image
        sheet.cell(row=row, column=5, value=item[list(item)[4]])
        # description_below_image
        sheet.cell(row=row, column=7, value=item[list(item)[6]])
        # classroom_exercises
        sheet.cell(row=row, column=8, value=item[list(item)[7]])
        # score
        sheet.cell(row=row, column=9, value=item[list(item)[8]])
        # number_of_errors
        sheet.cell(row=row, column=10, value=item[list(item)[9]])
        # category
        sheet.cell(row=row, column=11, value=CATEGORY_CHOICES[item[list(item)[10]] - 1][1])
        # fixed
        sheet.cell(row=row, column=12, value=item[list(item)[11]])
        # exam_name
        sheet.cell(row=row, column=13, value=item[list(item)[12]])
        # exam_times
        sheet.cell(row=row, column=14, value=item[list(item)[13]])
        # md5_value
        sheet.cell(row=row, column=15, value=item[list(item)[14]])
        # answer
        sheet.cell(row=row, column=16, value=item[list(item)[15]])
        # released
        sheet.cell(row=row, column=17, value=item[list(item)[16]])
        # creator
        sheet.cell(row=row, column=18, value=item[list(item)[17]])
        # created
        sheet.cell(row=row, column=19, value=item[list(item)[18]])
        # modified
        sheet.cell(row=row, column=20, value=item[list(item)[19]])

        # iamge
        if item[list(item)[5]]:
            img = PILImage.open(item[list(item)[5]])
            # img = img.resize((100, 100))  # Resize the image as needed

            img_io = BytesIO()
            img.save(img_io, 'PNG')

            img_path = f'media/study/images/tmp_image_{row}.png'
            with open(img_path, 'wb') as f:
                f.write(img_io.getvalue())

            img = Image(img_path)
            sheet.add_image(img, f'{IMAGE_COLUMN}{row}') # F means column=6

    for index, row in enumerate(sheet.iter_rows()):
        # Set row height, title no need to set
        if index != 0:
            sheet.row_dimensions[index + 1].height = 180
        else:
            sheet.row_dimensions[index + 1].height = 60
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical='center', horizontal='center')
                
    # Set row width
    sheet.column_dimensions['E'].width = 50 # description_above_image
    sheet.column_dimensions[IMAGE_COLUMN].width = 80 # image
    sheet.column_dimensions['G'].width = 50 # description_above_image
    # Save the Excel file
    workbook.save('media/study/files/output.xlsx')
    # Create a response with the Excel file
    response = HttpResponse(content_type="application/ms-excel")
    response["Content-Disposition"] = "attachment; filename=study_questions.xlsx"

    # Save the Excel workbook to the response
    workbook.save(response)

    return response


def generate_test_paper(request):
    # Fetch data from the Django model
    grade = request.GET.get('grade')
    subject = get_choice_key_by_value(SUBJECT_CHOICES, request.GET.get('subject'))
    creator = request.GET.get('creator', 'ben')
    is_random = True if request.GET.get('random') == u'true' else False
    if is_random:
        items = list(Question.objects.filter(grade=grade, subject=subject, released=True))
        size = int(request.GET.get('size')) if int(request.GET.get('size')) < len(items) else len(items)
        # random items
        selected_items = random.sample(items, size)
    else:
        items = Question.objects.filter(grade=grade, subject=subject, released=True).order_by('exam_times')
        size = int(request.GET.get('size')) if int(request.GET.get('size')) < len(items) else len(items)
        selected_items = items[:size]

    # Create a new Word document
    document = Document()
    exams = Exam.objects.filter(grade=grade, subject=subject)

    # Add a title to the document
    # import pdb; pdb.set_trace()
    name = arabic_numerals_to_chinese_numerals(grade) + \
                _('年级') + SUBJECT_CHOICES[int(subject) - 1][1] + _('试题') + \
                arabic_numerals_to_chinese_numerals(len(exams) + 1)
    heading = document.add_heading(name, 1)
    run = heading.runs[0]
    font = run.font
    font.bold = True  
    font.size = Pt(24)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set alignment to center
    document.add_paragraph()
    document.add_paragraph()

    # Add data from the model to the document
    total_score = 0
    for index, item in enumerate(selected_items):
        # Add item exam_times
        item.exam_times += 1
        item.save()
        # Add score
        total_score += item.score
        # Add question link
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        font = run.font
        font.size = Pt(12)   
        add_hyperlink(paragraph, f'({item.id})', f'http://mybackend.online:8000/admin/study/question/{item.id}/change/')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{arabic_numerals_to_chinese_numerals(index + 1)}、{item.title} ({item.score}{_('分')})")
        font = run.font
        font.name = '宋体'  # Set the font name
        font.size = Pt(16)    # Set the font size
        font.bold = True
        font.italic = False
        if item.description_above_image:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(item.get_description_above_image())
            font = run.font
            font.name = '宋体'  
            font.size = Pt(14)   
        if item.image:
            # Open the image using Pillow
            image = PILImage.open(item.image)
            # Get the original width and height of the image
            original_width, original_height = image.size
            width = 6
            i = original_width / width
            document.add_picture(item.image, width=Inches(width), height=Inches(original_height / i))

            # pic = document.add_picture("." + item.image.url, width=Inches(5), height=Inches(3))
            # pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if item.description_below_image:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(item.get_description_below_image())
            font = run.font
            font.name = '宋体'  
            font.size = Pt(14)  
        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()

    # Save exam
    if len(selected_items) > 0:
        exam = Exam.objects.create(
            grade=grade,
            subject=subject,
            name=name,
            total_score=total_score,
            creator=creator,
        )
        exam.questions.set(items)
        exam.save()

    # Create a response with the Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=my_data_export.docx'
    document.save(response)

    return response