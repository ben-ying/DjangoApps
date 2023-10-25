import random

from django.core.files.base import ContentFile
from django.http import HttpResponse
from django.http import HttpResponse
from django.utils.translation import gettext as _

from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from io import BytesIO
from PIL import Image as PILImage

from ..models import Question, Exam
from ..models import SUBJECT_CHOICES
from ..utils import get_choice_key_by_value, arabic_numerals_to_chinese_numerals, add_hyperlink


def generate_test_paper(request):
    # Fetch data from the Django model
    grade = request.GET.get('grade')
    subject = get_choice_key_by_value(SUBJECT_CHOICES, request.GET.get('subject'))
    print('subject: ' + request.GET.get('subject'))
    print('subject: ' + str(subject))
    creator = request.GET.get('creator', 'ben')
    is_random = True if request.GET.get('random') == u'true' else False
    if is_random:
        items = list(Question.objects.filter(grade=grade, subject=subject, released=True))
        size = int(request.GET.get('size')) if int(request.GET.get('size')) < len(items) else len(items)
        # random items
        selected_items = random.sample(items, size)
    else:
        items = Question.objects.filter(grade=grade, subject=subject, released=True).order_by('exam_times')
        size = int(request.GET.get('size')) if int(request.GET.get('size')) < len(items) else len(items)
        selected_items = items[:size]

    # Create a new Word document
    document = Document()
    exams = Exam.objects.filter(grade=grade, subject=subject)

    # Add a title to the document
    # import pdb; pdb.set_trace()
    name = arabic_numerals_to_chinese_numerals(grade) + \
                _('年级') + SUBJECT_CHOICES[int(subject) - 1][1] + _('试题') + \
                arabic_numerals_to_chinese_numerals(len(exams) + 1)
    heading = document.add_heading(name, 1)
    run = heading.runs[0]
    font = run.font
    font.bold = True  
    font.size = Pt(24)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set alignment to center
    document.add_paragraph()
    document.add_paragraph()

    # Add data from the model to the document
    total_score = 0
    for index, item in enumerate(selected_items):
        # Add item exam_times
        item.exam_times += 1
        item.save()
        # Add score
        total_score += item.score
        # Add question link
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        font = run.font
        font.size = Pt(12)   
        add_hyperlink(paragraph, f'({item.id})', f'http://mybackend.online:8000/admin/study/question/{item.id}/change/')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{arabic_numerals_to_chinese_numerals(index + 1)}、{item.title} ({item.score}{_('分')})")
        font = run.font
        font.name = '宋体'  # Set the font name
        font.size = Pt(16)    # Set the font size
        font.bold = True
        font.italic = False
        if item.description_above_image:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(item.get_description_above_image())
            font = run.font
            font.name = '宋体'  
            font.size = Pt(14)   
        if item.image:
            # Open the image using Pillow
            image = PILImage.open(item.image)
            # Get the original width and height of the image
            original_width, original_height = image.size
            width = 6
            i = original_width / width
            document.add_picture(item.image, width=Inches(width), height=Inches(original_height / i))

            # pic = document.add_picture("." + item.image.url, width=Inches(5), height=Inches(3))
            # pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if item.description_below_image:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(item.get_description_below_image())
            font = run.font
            font.name = '宋体'  
            font.size = Pt(14)  
        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()

    today = date.today()
    doc_name = f'{today}_{name}.docx'

    # Save exam
    if len(selected_items) > 0:
         # Save the document to a BytesIO object
        buffer = BytesIO()
        document.save(buffer)
        content = ContentFile(buffer.getvalue())
        buffer.close()

        exam = Exam.objects.create(
            grade=grade,
            subject=subject,
            name=name,
            total_score=total_score,
            creator=creator,
        )
        exam.questions.set(items)
        exam.document.save(doc_name, content, save=True)
        exam.save()

    # Create a response with the Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=exam.docx'
    document.save(response)

    return response
